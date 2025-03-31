import streamlit as st
import os
import pdfplumber
import docx2txt
import pandas as pd
from docx import Document

# Function to extract text from PDF
def extract_text_from_pdf(pdf_path):
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            extracted_text = page.extract_text()
            if extracted_text:
                text += extracted_text + "\n"
    return text

# Function to extract text from DOCX
def extract_text_from_docx(docx_path):
    return docx2txt.process(docx_path)

# Function to process uploaded file
def process_uploaded_file(file_path, file_type):
    if file_type == "pdf":
        extracted_text = extract_text_from_pdf(file_path)
    elif file_type == "docx":
        extracted_text = extract_text_from_docx(file_path)
    else:
        with open(file_path, "r", encoding="utf-8") as f:
            extracted_text = f.read()
    return extracted_text

# Function to classify requirements into functional & non-functional
def extract_requirements(text):
    functional_reqs = []
    non_functional_reqs = []
    
    functional_keywords = ["must", "should", "shall", "required", "implement", "feature"]
    non_functional_keywords = ["performance", "security", "usability", "scalability", "availability", "compliance"]

    lines = text.split("\n")
    for line in lines:
        if any(word in line.lower() for word in functional_keywords):
            functional_reqs.append(line.strip())
        elif any(word in line.lower() for word in non_functional_keywords):
            non_functional_reqs.append(line.strip())

    return functional_reqs, non_functional_reqs

# MoSCoW Prioritization
def prioritize_requirements(functional_reqs, non_functional_reqs):
    must_have = []
    should_have = []
    could_have = []
    wont_have = []

    for req in functional_reqs + non_functional_reqs:
        if "must" in req.lower() or "required" in req.lower():
            must_have.append(req)
        elif "should" in req.lower():
            should_have.append(req)
        elif "could" in req.lower():
            could_have.append(req)
        else:
            wont_have.append(req)

    return must_have, should_have, could_have, wont_have

# Function to save requirements to Excel
def save_to_excel(must_have, should_have, could_have, wont_have, output_path):
    df = pd.DataFrame({
        "MoSCoW Priority": ["Must Have"] * len(must_have) + ["Should Have"] * len(should_have) +
                           ["Could Have"] * len(could_have) + ["Won't Have"] * len(wont_have),
        "Requirement": must_have + should_have + could_have + wont_have
    })
    df.to_excel(output_path, index=False)

# Function to generate Word requirement document
def generate_requirement_document(must_have, should_have, could_have, wont_have, output_path):
    doc = Document()
    doc.add_heading("Requirement Document", level=1)

    for category, reqs in zip(["Must Have", "Should Have", "Could Have", "Won't Have"], [must_have, should_have, could_have, wont_have]):
        if reqs:
            doc.add_heading(category, level=2)
            for req in reqs:
                doc.add_paragraph(f"- {req}")

    doc.save(output_path)

# Streamlit UI
st.title("AI Requirement Gathering System")
st.write("Upload a document to extract requirements and prioritize them using MoSCoW.")

uploaded_file = st.file_uploader("Choose a file", type=["pdf", "docx", "txt"])

if uploaded_file is not None:
    # Save uploaded file
    save_dir = "uploads"
    os.makedirs(save_dir, exist_ok=True)
    file_path = os.path.join(save_dir, uploaded_file.name)

    with open(file_path, "wb") as f:
        f.write(uploaded_file.read())

    st.success(f"✅ File uploaded successfully: {uploaded_file.name}")

    # Extract text
    file_type = uploaded_file.name.split(".")[-1]
    extracted_text = process_uploaded_file(file_path, file_type)

    # Extract functional & non-functional requirements
    functional_reqs, non_functional_reqs = extract_requirements(extracted_text)

    # Prioritize requirements using MoSCoW
    must_have, should_have, could_have, wont_have = prioritize_requirements(functional_reqs, non_functional_reqs)

    # Display extracted requirements
    st.subheader("🔹 Functional Requirements")
    for req in functional_reqs:
        st.write(f"- {req}")

    st.subheader("🔹 Non-Functional Requirements")
    for req in non_functional_reqs:
        st.write(f"- {req}")

    # Display MoSCoW Prioritization
    st.subheader("📌 MoSCoW Prioritization")

    if must_have:
        st.write("✅ **Must Have** (Essential)")
        for req in must_have:
            st.write(f"- {req}")

    if should_have:
        st.write("🟡 **Should Have** (Important but not critical)")
        for req in should_have:
            st.write(f"- {req}")

    if could_have:
        st.write("🔵 **Could Have** (Nice to have)")
        for req in could_have:
            st.write(f"- {req}")

    if wont_have:
        st.write("❌ **Won’t Have** (Not needed)")
        for req in wont_have:
            st.write(f"- {req}")

    # Generate Excel file
    excel_output_path = os.path.join(save_dir, "requirements_moscow.xlsx")
    save_to_excel(must_have, should_have, could_have, wont_have, excel_output_path)
    st.download_button(label="📥 Download MoSCoW Prioritization (Excel)", data=open(excel_output_path, "rb"), file_name="requirements_moscow.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    # Generate Word document
    word_output_path = os.path.join(save_dir, "requirements_moscow.docx")
    generate_requirement_document(must_have, should_have, could_have, wont_have, word_output_path)
    st.download_button(label="📥 Download Structured Requirement Document", data=open(word_output_path, "rb"), file_name="requirements_moscow.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

else:
    st.info("Please upload a file to extract requirements.")
